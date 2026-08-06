"""마크다운 노트 → **Word(.docx)**.

왜 필요한가
-----------
이 도구가 내는 것은 마크다운이지만([D48](../../../docs/decisions.md#d48)),
증권사에서 리포트가 오가는 형식은 Word다. 초안을 애널리스트에게 넘길 때
([D51](../../../docs/decisions.md#d51)의 「넘김」) 마크다운 파일을 주면
받는 쪽이 다시 변환해야 한다.

**표를 진짜 Word 표로 만든다.** 리포트의 알맹이가 표에 있고, 파이프 문자가
그대로 박힌 문단을 주면 아무도 안 쓴다.

무엇을 옮기고 무엇을 버리는가
------------------------------
* 제목(`#`~`###`) → Word 제목 스타일. 목차·탐색이 붙는다
* 표 → Word 표 (머리행 굵게)
* 인용(`>`) → 들여쓴 문단. 미검증 레인 경고가 여기 있다
* 목록(`-`) → 글머리 기호
* **굵게** → 굵게
* 링크 `[글](주소)` → 글만 남기고 주소는 괄호로. Word 하이퍼링크는
  python-docx가 직접 못 만들고, 주소가 보이는 편이 인쇄본에서 낫다
"""

from __future__ import annotations

import io
import re

# 인라인 서식. 굵게만 처리한다 — 리서치 노트에서 기울임은 거의 안 쓰인다.
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_QUOTE = re.compile(r"^>\s?(.*)$")
_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")


def _cells(line: str) -> list[str]:
    """`| a | b |` → `["a", "b"]`. 이스케이프된 파이프는 셀 구분이 아니다."""
    body = line.strip().strip("|")
    parts = re.split(r"(?<!\\)\|", body)
    return [p.strip().replace("\\|", "|") for p in parts]


def _plain(text: str) -> str:
    """링크를 사람이 읽는 형태로. 주소는 괄호에 남긴다."""
    return _LINK.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)


def _write_runs(paragraph, text: str) -> None:
    """`**굵게**`를 실제 굵은 글씨로 넣는다."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(markdown: str, *, title: str = "") -> bytes:
    """마크다운 → .docx 바이트.

    **표를 진짜 표로 만든다.** 마크다운 표를 문단으로 흘려보내면 파이프가
    그대로 박혀 Word에서 쓸 수 없다.
    """
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    if title:
        doc.core_properties.title = title

    lines = _plain(markdown).split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 표 — 머리행 + 구분행 + 본문행
        if line.startswith("|") and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1].strip()):
            header = _cells(line)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_cells(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            # `cell.text = ""`를 먼저 넣지 않는다 — 빈 run이 하나 생겨서
            # 그 뒤에 붙인 굵은 글씨가 `runs[0]`이 아니게 된다. 새 셀의
            # 문단은 원래 비어 있다.
            for cell, text in zip(table.rows[0].cells, header, strict=False):
                cell.paragraphs[0].add_run(text).bold = True
            for row in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row, strict=False):
                    _write_runs(cell.paragraphs[0], text)
            doc.add_paragraph()
            continue

        if not line.strip():
            i += 1
            continue

        m = _HEADING.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        m = _QUOTE.match(line)
        if m:
            # 미검증 레인 경고가 여기 있다. 눈에 띄어야 한다.
            p = doc.add_paragraph(style="Intense Quote")
            _write_runs(p, m.group(1))
            i += 1
            continue

        m = _BULLET.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _write_runs(p, m.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _write_runs(p, line.strip())
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
