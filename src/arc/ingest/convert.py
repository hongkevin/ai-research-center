"""업로드 문서 → **마크다운**.

왜 필요한가
-----------
RA는 백지에서 시작하지 않는다. 커버 중인 종목이면 **자기가 쓴 노트가 이미
있다.** 그동안 「과거 리포트」는 데이터 확보가 막혀 보류돼 있었는데, 사실
우리가 구할 게 아니라 사용자가 갖고 있는 것이었다. 게다가 자기 문서라
저작권이 깨끗하고, 남의 리포트보다 값지다 — 자기 하우스 포맷, 자기 직전
추정, 자기 관점이다.

왜 마크다운인가
---------------
이 제품이 내는 것이 이미 마크다운이다([D36](../../../docs/decisions.md#d36)).
들어오는 것도 같은 형식이면 비교·편집·발간이 한 종류의 문서로 통일된다.
PDF를 PDF로 들고 있으면 아무것도 못 한다.

**업로드 문서의 숫자는 본문에 안 들어간다.**
--------------------------------------------
우리가 검산한 값이 아니다. 기사 레인([D45](../../../docs/decisions.md#d45))과
같은 취급이다 — 맥락과 **기준선**은 되지만 본문 수치의 출처는 못 된다.
그 선을 넘으면 「본문 숫자는 전부 레지스트리를 거친다」는 전제가 무너지고,
이 제품이 파는 것이 없어진다.

실측
----
코퍼스 PDF 200편 무작위 표본에서 **추출 실패 0건**. 「한글 0%」로 잡힌 2편은
깨진 게 아니라 영문 리포트였다. 한국 증권사 PDF의 폰트 CMap 문제는 일반적인
함정이 아니다.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

# 업로드 상한. 리서치 노트는 보통 10페이지 안쪽이고, 큰 것도 50페이지를 넘지
# 않는다. 상한이 없으면 한 번의 업로드가 워커를 오래 붙든다(워커 1개 전제).
MAX_BYTES = 20 * 1024 * 1024
MAX_PAGES = 80

TEXT_SUFFIXES = (".md", ".markdown", ".txt")


class ConvertError(Exception):
    """변환 실패. **화면에 그대로 보여줄 문장**이어야 한다."""


@dataclass
class Converted:
    """변환 결과."""

    markdown: str
    source_name: str
    kind: str  # pdf | docx | text
    pages: int = 0
    warnings: list[str] = field(default_factory=list)

    @property
    def chars(self) -> int:
        return len(self.markdown)


# 표 안의 셀이 줄바꿈으로 흩어져 오는 것을 줄인다. 세 칸 이상 연속 공백은
# 원문에서 열 구분이었을 가능성이 높다.
_COLUMNS = re.compile(r"[ \t]{3,}")
_BLANKS = re.compile(r"\n{3,}")


# 한국어 리서치 노트에 **실제로 나오는** 문자 영역. 여기 밖은 깨진 것이다.
_ALLOWED_RANGES = (
    (0x0000, 0x024F),  # ASCII · 라틴 확장 (영문 이름·티커)
    (0x2000, 0x206F),  # 일반 문장부호 (— … · ‘ ’)
    (0x20A0, 0x20CF),  # 통화기호 (₩ €)
    (0x2100, 0x214F),  # ™ № ℓ
    (0x2190, 0x22FF),  # 화살표 · 수학기호 (↑ ↓ ≒)
    (0x2460, 0x24FF),  # ① ②
    (0x25A0, 0x27BF),  # 도형 · 기타기호 (■ ● ★ ☞)
    (0x3000, 0x303F),  # CJK 문장부호 (「」 〈〉)
    (0x3131, 0x318F),  # 한글 낱자 (ㄱ ㅏ)
    (0x3200, 0x33FF),  # ㈜ ㎡ ℃
    (0x4E00, 0x9FFF),  # 한자
    (0xAC00, 0xD7A3),  # 한글 음절
    (0xF900, 0xFAFF),  # 한자 호환
    (0xFF00, 0xFFEF),  # 전각
)


def garbled_ratio(text: str) -> float:
    """**글자가 깨진 비율.** 폰트 CMap이 어긋나면 이렇게 나온다.

    실측: 같은 종목 리포트인데 한 편은 멀쩡하고(「1분기 영업이익은…」) 다른
    편은 `ۓΌ("ࣄۓ݌ľ ηԮࠒҀғ`였다. PDF가 자기 폰트에 글리프→유니코드
    대응표를 안 싣거나 잘못 실으면 추출기가 엉뚱한 코드포인트를 뱉는다.

    **허용 영역을 명시로 둔다.** 처음에는 「0x2500 미만은 통과」로 짰다가
    놓쳤다 — 깨진 글자가 아랍(U+06xx)·데바나가리(U+09xx)로 떨어지는데 그게
    전부 통과했다. 한국어 리서치 노트에 나올 수 있는 영역만 세어야 한다.

    **부분적으로만 깨진 문서는 못 잡는다** (2009년 리포트의 「매춗액」·「달성핚」
    — 깨진 글자가 여전히 정상 한글 음절이다). 그건 통계로 정상 문서와 안
    갈렸다. 그래서 변환 결과를 사람이 보고 넘기는 흐름이 필요하다.
    """
    solid = [c for c in text if not c.isspace()]
    if len(solid) < 200:
        return 0.0
    odd = sum(1 for c in solid if not any(lo <= ord(c) <= hi for lo, hi in _ALLOWED_RANGES))
    return odd / len(solid)


# 이보다 높으면 「깨진 것 같다」고 말한다. 정상 리포트 실측은 0.00~0.02다.
GARBLED_AT = 0.15


def _clean(text: str) -> str:
    text = text.replace("­", "").replace("\xa0", " ")
    text = _COLUMNS.sub(" · ", text)
    return _BLANKS.sub("\n\n", text).strip()


def _from_pdf(data: bytes, name: str) -> Converted:
    try:
        import fitz  # pymupdf
    except ImportError as exc:  # pragma: no cover - 배포 환경에서는 항상 있다
        raise ConvertError("PDF 변환기를 불러오지 못했습니다.") from exc

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ConvertError(f"PDF를 열지 못했습니다: {type(exc).__name__}") from exc

    warnings: list[str] = []
    pages = doc.page_count
    if pages > MAX_PAGES:
        warnings.append(f"{MAX_PAGES}쪽까지만 읽었습니다 (전체 {pages}쪽).")

    # **글자 크기로 제목을 가른다.** 본문보다 뚜렷하게 큰 줄이 제목이다.
    # 리서치 노트는 목차 구조가 얕아서 h2 하나로 충분하다.
    sizes: list[float] = []
    blocks_per_page: list[list[tuple[float, str]]] = []
    for page in doc[:MAX_PAGES]:
        lines: list[tuple[float, str]] = []
        for block in page.get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                size = max((s.get("size", 0.0) for s in spans), default=0.0)
                sizes.append(size)
                lines.append((size, text))
        blocks_per_page.append(lines)

    if not sizes:
        raise ConvertError(
            "이 PDF에서 글자를 찾지 못했습니다. 스캔 이미지로 만든 문서일 수 있습니다."
        )

    body = sorted(sizes)[len(sizes) // 2]  # 중앙값 = 본문 크기
    heading_at = body * 1.35

    out: list[str] = []
    for i, lines in enumerate(blocks_per_page, 1):
        if i > 1:
            out.append(f"\n<!-- {i}쪽 -->\n")
        for size, text in lines:
            if size >= heading_at and len(text) <= 60:
                out.append(f"\n## {text}\n")
            else:
                out.append(text)
    doc.close()

    md = _clean("\n".join(out))
    if len(md) < 100:
        warnings.append("글자가 거의 없습니다. 표나 그림 위주의 문서일 수 있습니다.")
    ratio = garbled_ratio(md)
    if ratio >= GARBLED_AT:
        warnings.append(
            f"글자가 깨져 나왔습니다 (이상 문자 {ratio * 100:.0f}%). "
            "이 PDF가 폰트 정보를 제대로 싣지 않았습니다 — 원본을 다른 방법으로 "
            "내보내거나, 텍스트를 직접 붙여넣어 주십시오."
        )
    return Converted(
        markdown=md, source_name=name, kind="pdf", pages=min(pages, MAX_PAGES), warnings=warnings
    )


def _from_docx(data: bytes, name: str) -> Converted:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ConvertError("Word 변환기를 불러오지 못했습니다.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ConvertError(f"Word 문서를 열지 못했습니다: {type(exc).__name__}") from exc

    out: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        # Heading 1~3만 제목으로. 그 아래는 본문과 구분이 흐릿하다.
        level = next((n for n in (1, 2, 3) if f"heading {n}" in style), 0)
        out.append(f"\n{'#' * (level + 1)} {text}\n" if level else text)

    # **표는 파이프 표로 옮긴다.** 리서치 노트의 알맹이가 표에 있다.
    for table in document.tables:
        rows = [[c.text.strip().replace("|", "\\|") for c in r.cells] for r in table.rows]
        rows = [r for r in rows if any(r)]
        if len(rows) < 2:
            continue
        out.append("")
        out.append("| " + " | ".join(rows[0]) + " |")
        out.append("|" + "---|" * len(rows[0]))
        for row in rows[1:]:
            out.append("| " + " | ".join(row) + " |")
        out.append("")

    md = _clean("\n".join(out))
    if not md:
        raise ConvertError("Word 문서에서 글자를 찾지 못했습니다.")
    return Converted(markdown=md, source_name=name, kind="docx")


def convert(data: bytes, filename: str) -> Converted:
    """업로드 바이트 → 마크다운. 실패하면 `ConvertError`.

    **한글(HWP)은 아직 못 읽는다.** 표준 파서가 없고 포맷이 두 종류(hwp/hwpx)라
    따로 붙여야 한다. 지금은 「PDF로 내보내서 올리십시오」라고 말한다 —
    못 하는 것을 되는 척하는 것보다 낫다.
    """
    if not data:
        raise ConvertError("빈 파일입니다.")
    if len(data) > MAX_BYTES:
        raise ConvertError(f"파일이 너무 큽니다 ({len(data) // 1024 // 1024}MB). 20MB까지 됩니다.")

    lower = filename.lower()
    if lower.endswith(".pdf") or data[:5] == b"%PDF-":
        return _from_pdf(data, filename)
    if lower.endswith(".docx"):
        return _from_docx(data, filename)
    if lower.endswith((".hwp", ".hwpx")):
        raise ConvertError("한글(HWP) 문서는 아직 읽지 못합니다. PDF로 내보낸 뒤 올려 주십시오.")
    if lower.endswith(".doc"):
        raise ConvertError(
            "구형 Word(.doc)는 읽지 못합니다. .docx나 PDF로 저장한 뒤 올려 주십시오."
        )
    if lower.endswith(TEXT_SUFFIXES):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = data.decode("cp949")  # 국내 문서는 아직 CP949가 남아 있다
            except UnicodeDecodeError as exc:
                raise ConvertError("글자 인코딩을 알 수 없습니다.") from exc
        return Converted(markdown=_clean(text), source_name=filename, kind="text")

    raise ConvertError(f"지원하지 않는 형식입니다: {filename}. PDF · DOCX · MD를 받습니다.")
