"""마크다운 노트 → Word.

증권사에서 리포트가 오가는 형식은 Word다. 초안을 애널리스트에게 넘길 때
(D51의 「넘김」) 마크다운을 주면 받는 쪽이 다시 변환해야 한다.

**표가 진짜 표여야 한다.** 리포트의 알맹이가 표에 있고, 파이프 문자가 그대로
박힌 문단을 주면 아무도 안 쓴다.
"""

from __future__ import annotations

import io

import docx

from arc.render.docx import markdown_to_docx

NOTE = """# 삼성물산(주) (028260) — 2025년 연간 실적 리뷰

| 항목 | 내용 |
|---|---|
| 시장 | KOSPI |
| 감사의견 | 적정의견 · 삼일회계법인 |

## 0. 공시 밖 배경

> ⚠ **공시 밖 · 검증 필요** — 이 절은 사업보고서에 근거가 없습니다.

### 0.1 산업 구조

건설과 상사가 **섞인** 구조다.

- 첫째 항목
- 둘째 항목

자세한 것은 [공시 원문](https://dart.fss.or.kr/x)에 있다.
"""


def _doc(md: str = NOTE):
    return docx.Document(io.BytesIO(markdown_to_docx(md)))


class TestStructure:
    def test_headings_become_word_headings(self):
        """Word 제목 스타일이어야 목차·탐색이 붙는다."""
        d = _doc()
        styles = {p.style.name for p in d.paragraphs if p.text.strip()}
        assert "Heading 1" in styles and "Heading 2" in styles and "Heading 3" in styles

    def test_tables_are_real_tables(self):
        """**파이프가 박힌 문단을 주면 아무도 안 쓴다.**"""
        d = _doc()
        assert len(d.tables) == 1
        t = d.tables[0]
        assert [c.text for c in t.rows[0].cells] == ["항목", "내용"]
        assert [c.text for c in t.rows[1].cells] == ["시장", "KOSPI"]
        assert not any("|" in p.text for p in d.paragraphs)

    def test_table_header_is_bold(self):
        t = _doc().tables[0]
        assert t.rows[0].cells[0].paragraphs[0].runs[0].bold is True

    def test_quote_keeps_the_unverified_warning(self):
        """미검증 레인 경고는 눈에 띄어야 한다 (D31)."""
        d = _doc()
        quotes = [p for p in d.paragraphs if p.style.name == "Intense Quote"]
        assert quotes and "공시 밖" in quotes[0].text

    def test_bullets_become_bullets(self):
        d = _doc()
        bullets = [p for p in d.paragraphs if p.style.name == "List Bullet"]
        assert [p.text for p in bullets] == ["첫째 항목", "둘째 항목"]


class TestInline:
    def test_bold_survives(self):
        d = _doc()
        p = next(p for p in d.paragraphs if p.text.startswith("건설과 상사가"))
        assert any(r.bold and r.text == "섞인" for r in p.runs)
        assert "**" not in p.text

    def test_links_keep_the_address(self):
        """인쇄본에서는 주소가 보이는 편이 낫다."""
        d = _doc()
        text = "\n".join(p.text for p in d.paragraphs)
        assert "공시 원문 (https://dart.fss.or.kr/x)" in text
        assert "](" not in text

    def test_escaped_pipe_is_not_a_column(self):
        """산식의 `\\|전기\\|`가 셀을 밀어내면 안 된다 (전에 한 번 밟았다)."""
        md = "| 산식 | 값 |\n|---|---|\n| (당기 - 전기) \\| 절대값 | 3 |\n"
        t = _doc(md).tables[0]
        assert len(t.rows[1].cells) == 2
        assert "|" in t.rows[1].cells[0].text


class TestEdges:
    def test_empty_note_does_not_raise(self):
        assert len(markdown_to_docx("")) > 0

    def test_table_without_body_rows(self):
        d = _doc("| a | b |\n|---|---|\n")
        assert len(d.tables) == 1
