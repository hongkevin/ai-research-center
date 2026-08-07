"""업로드 문서 → 마크다운, 그리고 직전 노트 읽기.

RA는 백지에서 시작하지 않는다. 커버 중인 종목이면 자기가 쓴 노트가 이미 있고,
그게 우리가 만들 초안의 **기준선**이자 **형식**이다.

지켜야 할 것: **업로드 문서의 숫자는 본문에 안 들어간다.** 우리가 검산한 값이
아니다. 비교 패널과 구성 힌트로만 쓴다.
"""

from __future__ import annotations

import io
import zipfile

import pytest

from arc.ingest.convert import GARBLED_AT, ConvertError, convert, garbled_ratio
from arc.ingest.prior import outline_of, parse_extraction, read_prior


class TestConvertGuards:
    def test_empty_file(self):
        with pytest.raises(ConvertError, match="빈 파일"):
            convert(b"", "x.pdf")

    def test_too_big(self):
        with pytest.raises(ConvertError, match="너무 큽니다"):
            convert(b"x" * (21 * 1024 * 1024), "x.pdf")

    def test_hwp_says_what_to_do_instead(self):
        """**못 하는 것을 되는 척하지 않는다.** 대신 할 일을 알려준다."""
        with pytest.raises(ConvertError, match="PDF로 내보낸 뒤"):
            convert(b"\xd0\xcf\x11\xe0", "리포트.hwp")

    def test_old_word_is_refused_with_a_way_out(self):
        with pytest.raises(ConvertError, match=r"\.docx"):
            convert(b"\xd0\xcf\x11\xe0", "리포트.doc")

    def test_unknown_suffix_lists_what_we_take(self):
        with pytest.raises(ConvertError, match="PDF · DOCX · MD"):
            convert(b"data", "리포트.pptx")


class TestConvertText:
    def test_markdown_passes_through(self):
        got = convert("## 제목\n\n본문입니다.".encode(), "note.md")
        assert got.kind == "text"
        assert "## 제목" in got.markdown

    def test_cp949_is_decoded(self):
        """국내 문서에는 CP949가 아직 남아 있다."""
        got = convert("삼성물산 실적 리뷰".encode("cp949"), "note.txt")
        assert "삼성물산" in got.markdown

    def test_wide_gaps_become_separators(self):
        """PDF·텍스트 표는 열이 공백으로 벌어져 온다."""
        got = convert(b"\xea\xb0\x80    \xeb\x82\x98", "t.txt")
        assert "·" in got.markdown


class TestGarbled:
    def test_clean_korean_scores_zero(self):
        text = "1분기 영업이익은 전년 동기 대비 68% 증가한 381억원을 기록했다. " * 12
        assert garbled_ratio(text) < 0.01

    def test_broken_cmap_is_caught(self):
        """실측 형태 — 깨진 글자가 아랍·데바나가리 영역으로 떨어진다.

        한때 「0x2500 미만은 통과」로 짜서 이걸 통째로 놓쳤다.
        """
        text = "ۓΌ(ण2ࣄۓ݌ľ ηԮࠒҀғ ۵ϊ ⰲⴙᤱⷝ㒍パ㓅㊆ " * 20
        assert garbled_ratio(text) >= GARBLED_AT

    def test_english_report_is_not_garbled(self):
        """영문 리포트는 깨진 게 아니다 — 실측에서 오탐이 났던 자리."""
        text = "Maintain Hold, with TP range of 70,000-99,000. Our outlook is good. " * 8
        assert garbled_ratio(text) < 0.01

    def test_short_text_is_not_judged(self):
        """표본이 적으면 판정하지 않는다. 거짓 경고가 없는 것보다 나쁘다."""
        assert garbled_ratio("ۓΌ") == 0.0


class TestOutline:
    def test_reads_the_section_order(self):
        md = "## 투자포인트\n본문\n### 1. 실적\n본문\n## 밸류에이션\n"
        assert outline_of(md) == ["투자포인트", "1. 실적", "밸류에이션"]

    def test_skips_page_markers_and_repeats(self):
        md = "## 요약\n<!-- 2쪽 -->\n## 요약\n## 리스크\n"
        assert outline_of(md) == ["요약", "리스크"]

    def test_no_headings_is_empty_not_an_error(self):
        assert outline_of("본문만 있는 문서") == []


class TestParseExtraction:
    def test_reads_target_and_estimates(self):
        target, rating, rows = parse_extraction(
            {
                "target_price": 21000,
                "rating": "Hold",
                "estimates": [
                    {"year": 2026, "revenue": 10373000000000, "operating_income": 1133000000000}
                ],
            }
        )
        assert (target, rating) == (21000, "Hold")
        assert rows[0]["year"] == 2026 and rows[0]["revenue"] == 10373000000000

    def test_drops_rows_whose_year_is_not_a_year(self):
        """표를 잘못 읽으면 연도 칸에 금액이 들어온다."""
        _, _, rows = parse_extraction({"estimates": [{"year": 10373, "revenue": 1}]})
        assert rows == []

    def test_drops_rows_with_no_values(self):
        _, _, rows = parse_extraction({"estimates": [{"year": 2026}]})
        assert rows == []

    def test_strips_units_from_strings(self):
        _, _, rows = parse_extraction(
            {"target_price": "21,000원", "estimates": [{"year": "2026", "revenue": "1,000"}]}
        )
        assert rows[0]["revenue"] == 1000

    def test_sorted_by_year(self):
        _, _, rows = parse_extraction(
            {"estimates": [{"year": 2027, "revenue": 2}, {"year": 2025, "revenue": 1}]}
        )
        assert [r["year"] for r in rows] == [2025, 2027]


class TestReadPrior:
    def test_outline_works_without_an_llm(self):
        """**차례만 있어도 「구성 따라 쓰기」는 성립한다.**"""
        note = read_prior(None, "## 투자포인트\n## 밸류에이션\n", "prior.pdf")
        assert note.outline == ["투자포인트", "밸류에이션"]
        assert note.usable
        assert note.problems

    def test_llm_failure_keeps_the_outline(self):
        class Boom:
            def complete(self, **_):
                raise RuntimeError("down")

        note = read_prior(Boom(), "## 요약\n", "prior.pdf")
        assert note.outline == ["요약"]
        assert any("RuntimeError" in p for p in note.problems)

    def test_reads_estimates_from_a_fenced_reply(self):
        """모델이 ```json 울타리를 자주 씌운다."""

        class Stub:
            def complete(self, **_):
                class C:
                    text = (
                        '```json\n{"target_price": 21000, "rating": "Hold", '
                        '"estimates": [{"year": 2026, "revenue": 1000}]}\n```'
                    )

                return C()

        note = read_prior(Stub(), "## 요약\n", "prior.pdf")
        assert note.target_price == 21000
        assert note.estimates == [{"year": 2026, "revenue": 1000}]


def _docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """최소 DOCX 하나. python-docx로 만든다."""
    import docx

    d = docx.Document()
    for style, text in paragraphs:
        d.add_paragraph(text, style=style) if style else d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestConvertDocx:
    def test_headings_become_markdown_headings(self):
        data = _docx([("Heading 1", "투자포인트"), ("", "본문 문장입니다.")])
        got = convert(data, "note.docx")
        assert got.kind == "docx"
        assert "## 투자포인트" in got.markdown
        assert "본문 문장입니다." in got.markdown

    def test_broken_docx_is_reported_not_raised_raw(self):
        bad = io.BytesIO()
        with zipfile.ZipFile(bad, "w") as z:
            z.writestr("junk.txt", "not a docx")
        with pytest.raises(ConvertError, match="열지 못했습니다"):
            convert(bad.getvalue(), "note.docx")


class TestCompareAreas:
    """영역별 대조 (D64).

    숫자 비교(D46)는 「얼마가 달라졌나」를 말한다. 이쪽은 **「무엇이
    달라졌나」** — RA가 새 분기 노트를 열고 처음 묻는 것이 그쪽이다.
    """

    def _stub(self, payload):
        import json as _json

        class Stub:
            def complete(self, **_):
                class C:
                    text = _json.dumps(payload, ensure_ascii=False)

                return C()

        return Stub()

    def test_areas_are_read(self):
        from arc.ingest.prior import compare_areas

        areas, problems = compare_areas(
            self._stub(
                {
                    "areas": [
                        {
                            "area": "수익성",
                            "prior": "마진 개선을 전망했다",
                            "now": "이번 공시에서 마진이 악화됐다",
                            "verdict": "약화",
                        }
                    ]
                }
            ),
            "## 리포트\n본문",
            ["영업이익률이 나빠졌다"],
        )
        assert problems == []
        assert areas[0].area == "수익성" and areas[0].verdict == "약화"

    def test_numbers_get_the_area_dropped(self):
        """**직전 리포트의 숫자는 우리가 검산하지 않았다** — 섞이면 안 된다."""
        from arc.ingest.prior import compare_areas

        areas, problems = compare_areas(
            self._stub(
                {
                    "areas": [
                        {
                            "area": "실적",
                            "prior": "매출 12조를 전망했다",
                            "now": "유지됐다",
                            "verdict": "유지",
                        }
                    ]
                }
            ),
            "## 리포트",
            [],
        )
        assert areas == []
        assert problems and "숫자" in problems[0]

    def test_unknown_verdict_becomes_uncertain(self):
        """**확인불가를 두려워하지 않는다.** 공시 밖 주장은 그렇게 표시한다."""
        from arc.ingest.prior import compare_areas

        areas, _ = compare_areas(
            self._stub(
                {"areas": [{"area": "리스크", "prior": "가", "now": "나", "verdict": "몰라"}]}
            ),
            "## 리포트",
            [],
        )
        assert areas[0].verdict == "확인불가"

    def test_incomplete_area_is_skipped(self):
        from arc.ingest.prior import compare_areas

        areas, problems = compare_areas(
            self._stub({"areas": [{"area": "수익성", "prior": "", "now": "나"}]}),
            "## 리포트",
            [],
        )
        assert areas == [] and problems

    def test_no_llm_no_areas(self):
        from arc.ingest.prior import compare_areas

        areas, problems = compare_areas(None, "## 리포트", [])
        assert areas == [] and problems

    def test_provider_failure_does_not_raise(self):
        from arc.ingest.prior import compare_areas

        class Boom:
            def complete(self, **_):
                raise RuntimeError("down")

        areas, problems = compare_areas(Boom(), "## 리포트", [])
        assert areas == [] and any("RuntimeError" in p for p in problems)

    def test_prior_numbers_are_masked_before_the_prompt(self):
        """스니펫과 같은 규칙 — 유혹을 애초에 없앤다 (D45)."""
        import json as _json

        seen = {}

        class Spy:
            def complete(self, *, system, user, tier=None):
                seen["user"] = user

                class C:
                    text = _json.dumps({"areas": []})

                return C()

        from arc.ingest.prior import compare_areas

        compare_areas(Spy(), "매출액은 12조 3,456억원이다.", [])
        assert "12조 3,456억원" not in seen["user"]
